# export/docx_exporter.py
from docx import Document
from docx.shared import Pt
import os

def export_to_docx(passage, question_blocks, output_path):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # 지문
    doc.add_heading('【지문】', level=1)
    para = doc.add_paragraph(passage)
    para.style.font.size = Pt(12)

    # 문제
    doc.add_heading('【문제】', level=1)
    for i, block in enumerate(question_blocks, 1):
        q_text = block[0].split(". ", 1)[1] if ". " in block[0] else block[0]
        p = doc.add_paragraph(f"{i}. {q_text}", style='List Number')
        p.style.font.size = Pt(12)
        for line in block[1:]:
            sub = doc.add_paragraph(line)
            sub.style.font.size = Pt(12)

    doc.save(output_path)
    print(f"[INFO] 문제지가 저장되었습니다: {output_path}")